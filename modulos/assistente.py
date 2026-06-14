import streamlit as st
import pandas as pd
import google.generativeai as genai
import io
import os
from docx import Document

# ==========================================
# 1. WIDGET DO CHAT DO BENTO (Flutuante / Tela Cheia)
# ==========================================
def render_widget_ia(supabase):
    is_master = (st.session_state.get('perfil_logado') == "Master") or (st.session_state.get('usuario_logado') in ['breno', 'uriel'])

    if "GEMINI_API_KEY" not in st.secrets:
        st.sidebar.warning("⚠️ Chave API da IA não configurada.")
        return

    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

    if "mensagens_ia" not in st.session_state:
        st.session_state["mensagens_ia"] = [{"role": "assistant", "content": "Olá! Sou o Bento, a Inteligência Artificial da Consorbens. Qual a sua dúvida sobre as administradoras?"}]

    if "bento_tela_cheia" not in st.session_state:
        st.session_state["bento_tela_cheia"] = False

    # --- A MÁGICA DO CSS PARA BLINDAR O SCROLL ---
    # Bloqueamos o scroll da janela principal para que os botões do topo NUNCA sumam.
    css_popover = """
    <style>
        div[data-testid="stPopoverBody"] {
            overflow-y: hidden !important; /* Trava o scroll da janela externa */
            padding: 1rem !important;
        }
    """
    
    if st.session_state["bento_tela_cheia"]:
        css_popover += """
            div[data-testid="stPopoverBody"] {
                position: fixed !important;
                top: 5vh !important;
                left: 5vw !important;
                width: 90vw !important;
                min-width: 90vw !important;
                height: 90vh !important;
                max-height: 90vh !important;
                z-index: 99999 !important;
                border-radius: 12px !important;
                box-shadow: 0 10px 40px rgba(0,0,0,0.4) !important;
            }
        """
    css_popover += "</style>"
    st.markdown(css_popover, unsafe_allow_html=True)

    st.sidebar.divider()
    
    with st.sidebar.popover("🤖 Falar com o Bento", use_container_width=True):
        
        # Cabeçalho SUPER COMPACTO e fixo no topo
        c1, c2, c3 = st.columns([5, 2.5, 2.5])
        with c1:
            st.markdown("#### 🤖 Bento AI")
        with c2:
            if st.button("🗑️ Limpar", help="Apagar histórico", use_container_width=True):
                st.session_state["mensagens_ia"] = [{"role": "assistant", "content": "Olá! Sou o Bento, a Inteligência Artificial da Consorbens. Qual a sua dúvida sobre as administradoras?"}]
                st.rerun()
        with c3:
            if st.session_state["bento_tela_cheia"]:
                if st.button("↙️ Voltar", help="Tamanho original", use_container_width=True):
                    st.session_state["bento_tela_cheia"] = False
                    st.rerun()
            else:
                if st.button("⛶ Ampliar", help="Tela Cheia", use_container_width=True):
                    st.session_state["bento_tela_cheia"] = True
                    st.rerun()

        # Formulário de Pergunta
        with st.form("chat_form", clear_on_submit=True):
            c_input, c_btn = st.columns([5, 1])
            pergunta = c_input.text_input("Dúvida:", label_visibility="collapsed", placeholder="Ex: Qual a taxa da Yamaha?")
            enviou = c_btn.form_submit_button("➤", use_container_width=True)
            
            if enviou and pergunta:
                st.session_state["mensagens_ia"].append({"role": "user", "content": pergunta})
                
                with st.spinner("O Bento está pensando..."):
                    try:
                        # Auto-Detector de Modelos
                        modelos_permitidos = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
                        
                        if not modelos_permitidos:
                            st.error("⚠️ O Google bloqueou sua chave. Crie um novo projeto no Google Cloud escolhendo os Estados Unidos (US).")
                        else:
                            modelo_escolhido = next((m for m in modelos_permitidos if '1.5-flash' in m), modelos_permitidos[0])
                            modelo_limpo = modelo_escolhido.replace('models/', '')
                            
                            res = supabase.table("base_conhecimento_ia").select("*").execute()
                            df_base = pd.DataFrame(res.data)
                            
                            contexto_geral = ""
                            if not df_base.empty:
                                for _, row in df_base.iterrows():
                                    contexto_geral += f"\n\n--- ADMINISTRADORA: {row['administradora']} ---\n"
                                    contexto_geral += f"Regras Operacionais: {row['regras_operacionais']}\n"
                                    if is_master:
                                        contexto_geral += f"Comissões (SIGILOSO): {row['regras_comissionamento']}\n"
                            
                            if not contexto_geral:
                                contexto_geral = "Ainda não há regras cadastradas no banco de dados."

                            model = genai.GenerativeModel(modelo_limpo)
                            
                            prompt_sistema = f"""Você é o Bento, o assistente virtual e mascote exclusivo do ERP da Consorbens, especializado em consórcios.
                            Sua personalidade é extremamente prestativa, simpática e profissional.
                            
                            Responda APENAS com base neste contexto (se não souber, diga que não tem a informação):
                            {contexto_geral}
                            
                            PERGUNTA DO USUÁRIO: {pergunta}
                            """
                            
                            resposta = model.generate_content(prompt_sistema)
                            st.session_state["mensagens_ia"].append({"role": "assistant", "content": resposta.text})
                    except Exception as e:
                        st.error(f"Erro ao consultar o Bento: {e}")

        # Ajuste dinâmico da altura da caixa para caber perfeitamente na tela
        altura_caixa = 480 if st.session_state["bento_tela_cheia"] else 350
        chat_container = st.container(height=altura_caixa)
        
        with chat_container:
            for msg in st.session_state["mensagens_ia"]:
                if msg["role"] == "assistant":
                    st.chat_message(msg["role"], avatar="🤖").write(msg["content"])
                else:
                    st.chat_message(msg["role"]).write(msg["content"])

# ==========================================
# 2. TELA DE CADASTRO DE REGRAS (Só para Masters)
# ==========================================
def render_config_ia(supabase):
    st.markdown("### ⚙️ Base de Conhecimento do Bento")
    st.markdown("Aqui você escreve tudo o que o Bento precisa saber para responder à equipe com precisão.")
    
    try:
        res = supabase.table("base_conhecimento_ia").select("*").execute()
        df_bd = pd.DataFrame(res.data)
    except:
        df_bd = pd.DataFrame(columns=["id", "administradora", "regras_operacionais", "regras_comissionamento"])
        
    with st.expander("➕ Cadastrar Regras Manualmente", expanded=False):
        with st.form("form_ia_cadastro"):
            admin_nome = st.text_input("Nome da Administradora *")
            reg_op = st.text_area("Regras Operacionais (Visível para todos)", placeholder="Idade do veículo, lances embutidos, prazos...", height=150)
            reg_com = st.text_area("Regras de Comissionamento (SIGILOSO - Só Masters)", placeholder="Taxas da corretora, parcelamento de bônus...", height=150)
            
            if st.form_submit_button("Salvar na Base da IA", type="primary"):
                if admin_nome:
                    supabase.table("base_conhecimento_ia").insert({
                        "administradora": admin_nome,
                        "regras_operacionais": reg_op,
                        "regras_comissionamento": reg_com
                    }).execute()
                    st.success("Salvo com sucesso!")
                    st.rerun()
                else:
                    st.error("Preencha o nome da administradora.")
                    
    st.divider()

    st.markdown("#### 📥 Importar / 📤 Exportar Base via Word (.docx)")
    c_imp, c_exp = st.columns(2)
    
    with c_imp:
        st.write("**Importar arquivo do Word**")
        st.caption("O arquivo deve seguir a estrutura gerada no botão de exportação ao lado.")
        uploaded_file = st.file_uploader("Subir documento Word (.docx)", type=['docx'], label_visibility="collapsed")
        
        if uploaded_file is not None:
            if st.button("Processar e Salvar Documento Word", type="primary"):
                try:
                    doc = Document(uploaded_file)
                    current_admin = None
                    current_section = None
                    records = {}
                    
                    for p in doc.paragraphs:
                        text = p.text.strip()
                        if not text:
                            if current_admin and current_section:
                                records[current_admin][current_section] += "\n"
                            continue
                        
                        if text.upper().startswith("ADMINISTRADORA:"):
                            current_admin = text.split(":", 1)[1].strip().upper()
                            records[current_admin] = {"op": "", "com": ""}
                            current_section = None
                        elif text.upper() == "[REGRAS OPERACIONAIS]":
                            current_section = "op"
                        elif text.upper() == "[REGRAS DE COMISSIONAMENTO]":
                            current_section = "com"
                        elif text.startswith("--------------------"):
                            current_section = None
                        else:
                            if current_admin and current_section:
                                records[current_admin][current_section] += text + "\n"
                    
                    contador = 0
                    for admin, info in records.items():
                        if admin:
                            op_clean = info["op"].strip()
                            com_clean = info["com"].strip()
                            
                            existe = supabase.table("base_conhecimento_ia").select("id").eq("administradora", admin).execute()
                            
                            payload = {
                                "administradora": admin,
                                "regras_operacionais": op_clean,
                                "regras_comissionamento": com_clean
                            }
                            
                            if existe.data:
                                supabase.table("base_conhecimento_ia").update(payload).eq("id", existe.data[0]['id']).execute()
                            else:
                                supabase.table("base_conhecimento_ia").insert(payload).execute()
                            contador += 1
                            
                    st.success(f"✅ Sucesso! {contador} administradora(s) processada(s) e salvas no banco!")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Erro ao ler o arquivo Word: {e}")
                    
    with c_exp:
        st.write("**Exportar Base Atual para Word**")
        st.caption("Baixe um documento estruturado pronto para editar ou servir de modelo.")
        
        if not df_bd.empty:
            doc_out = Document()
            doc_out.add_heading('Base de Conhecimento - Consorbens IA', 0)
            p_inst = doc_out.add_paragraph()
            p_inst.add_run("INSTRUÇÕES DE EDIÇÃO:\n").bold = True
            p_inst.add_run("Você pode alterar os textos abaixo ou criar novas administradoras seguindo o mesmo padrão de tags corporativas em letras maiúsculas. Não altere as palavras entre colchetes.")
            
            for _, row in df_bd.iterrows():
                doc_out.add_paragraph(f"ADMINISTRADORA: {row['administradora'].upper()}")
                doc_out.add_paragraph("[REGRAS OPERACIONAIS]")
                doc_out.add_paragraph(row['regras_operacionais'] if row['regras_operacionais'] else "Nenhuma cadastrada")
                doc_out.add_paragraph("[REGRAS DE COMISSIONAMENTO]")
                doc_out.add_paragraph(row['regras_comissionamento'] if row['regras_comissionamento'] else "Nenhuma cadastrada")
                doc_out.add_paragraph("-" * 50)
                
            bio = io.BytesIO()
            doc_out.save(bio)
            
            st.download_button(
                label="⬇️ Baixar Base Completa em Word (.docx)",
                data=bio.getvalue(),
                file_name=f"base_conhecimento_ia_{pd.Timestamp.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.info("Cadastre pelo menos uma regra para habilitar o download.")

    st.divider()

    if not df_bd.empty:
        st.markdown("#### 📚 Base Atual de Conhecimento")
        for _, row in df_bd.iterrows():
            with st.expander(f"Administradora: {row['administradora']}"):
                st.markdown("**Regras Operacionais:**")
                st.write(row['regras_operacionais'])
                st.markdown("**Regras de Comissionamento:**")
                st.write(row['regras_comissionamento'])
                if st.button(f"🚨 Apagar Regras da {row['administradora']}", key=f"del_{row['id']}"):
                    supabase.table("base_conhecimento_ia").delete().eq("id", row['id']).execute()
                    st.rerun()
